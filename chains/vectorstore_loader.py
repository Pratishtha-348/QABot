import os
import requests
from PyPDF2 import PdfReader
import docx
from bs4 import BeautifulSoup
from langchain.docstore.document import Document
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings


# 🔹 Shared helper for vectorstore
def _build_vectorstore(text, persist_directory):
    docs = [Document(page_content=text)]
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    return Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory=persist_directory
    )


# 🔹 Load from uploaded file
def load_vectorstore(uploaded_file, persist_directory="chroma_dbs/file"):
    ext = os.path.splitext(uploaded_file.name)[1].lower()

    if ext == ".pdf":
        reader = PdfReader(uploaded_file)
        text = "\n".join([page.extract_text() or "" for page in reader.pages])

    elif ext == ".txt":
        text = uploaded_file.read().decode("utf-8")

    elif ext in [".docx", ".doc"]:
        doc = docx.Document(uploaded_file)
        text = "\n".join([p.text for p in doc.paragraphs])

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return _build_vectorstore(text, persist_directory)


# 🔹 Load from URL
def load_vectorstore_from_url(url, persist_directory="chroma_dbs/url"):
    response = requests.get(url)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    # Extract readable text from <p>, <li>, and headings
    text = "\n".join([t.get_text() for t in soup.find_all(["p", "li", "h1", "h2", "h3"])])

    return _build_vectorstore(text, persist_directory)
